# *********************************************************************
# 内容描述：  智能试卷题库管理系统
#           基于 Python/CustomTkinter 开发，包含题库增删改查、智能题型识别、统计分析、Word导入导出等功能。
# 创建人：
# 创建日期:  2025-12-22
# 修改历史： [一、初始版：2025-12-22,
# 1. 新增多选功能，支持批量删除题目；
# 2. 修复英文题库导入时选项 / 答案错误筛入题目栏的问题；
# 3. 新增管理员账号密码登录功能；
# 4. 新增搜索框，支持过滤表格题库内容；
# 5. 新增文件夹式题库层级界面，支持题库保存；
# 6. 新增题型 / 答案自动识别功能（规则推断题型、扩展 TXT 解析逻辑）；
# 7. 优化导出功能，支持 Word/PDF 格式（替换 JSON）；
# 8. 新增数据分析可视化功能，添加 “统计” 按钮并生成题型分布 / 导入历史饼图；
# 9. 替换存储引擎为 sqlite3（替代 JSON 读写）；
# 10. 基于 CustomTkinter 美化 UI（替换原生 Tkinter）；
#二、第二版：2025-12-22,
# 1. 实现选项和答案自动识别；
# 2. 标注多选删除题目功能未实现；
# 3. 新增搜索后返回数据库页面的功能；
# 4. 优化字体显示，修复题目 ID / 题型 / 内容 / 选项 / 答案字体过小问题；
# 5. 优化导出逻辑，支持按题目顺序导出、选择部分题目导出。
# 三、第三版：2025-12-22,
# 1. 新增 ID 左侧多选框，支持勾选删除 / 导入 / 导出题目；
# 2. 新增字体自定义功能（Ctrl + 滚轮缩放）；
# 3. 完善自动识别能力，实现题型 / 选项 / 答案全量自动推断；
# 4. 修复 UI 界面退出登录按钮显示问题，修改界面主题；
# 5. 新增 ID 排序按钮，支持 ID 升序 / 降序排序；
# 6. 修复统计饼状图数字与题型中文重叠问题。
# 四、第四版：2025-12-23,
# 1. 新增 “保存全部导出” 按钮；
# 2. 优化选中交互，单击题目行显示蓝色底色标识选中状态。
# 五、第五版：2025-12-23,
# 1. 优化统计饼状图，合并中英文含义相同的题型分类；
# 2. 新增饼状图交互功能（Ctrl + 滚轮缩放）。]
# 遗留问题： 暂无
# *********************************************************************

import customtkinter as ctk
from tkinter import ttk, messagebox, filedialog
import sqlite3
import os
import re
from docx import Document
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from collections import defaultdict

# --- 全局配置 ---
# 设置外观模式为跟随系统 (System: 跟随系统, Light: 亮色, Dark: 暗色)
ctk.set_appearance_mode("System")
# 设置默认颜色主题为蓝色 (按钮高亮色、进度条颜色等)
ctk.set_default_color_theme("blue")


# *********************************************************************
# 类名：StringResources
# 功能描述：静态资源类。
#           将所有的字符串常量（标题、按钮文字、提示信息）集中管理。
#           优点：方便后续修改文案或进行多语言国际化，避免代码中出现"魔法字符串"。
# *********************************************************************
class StringResources:
    TITLE_MAIN = "智能试卷题库管理系统"
    TITLE_LOGIN = "🔒 系统登录"
    TITLE_STATS = "统计分布"
    TITLE_PREVIEW = "导入预览"

    BTN_LOGIN = "登 录"
    BTN_EXIT = "退出登录"
    BTN_ADD_BANK = "➕ 新建题库"
    BTN_DEL_BANK = "❌ 删除题库"
    BTN_ADD_Q = "➕ 录入"
    BTN_DEL_CHECKED = "🗑️ 删除勾选"
    BTN_STATS = "📊 统计分布"
    BTN_IMPORT = "📥 导入"
    BTN_EXPORT_ALL = "📤 导出全部"
    BTN_EXPORT_CHECKED = "📤 导出勾选"
    BTN_SEARCH = "搜索"
    BTN_RESET = "重置"

    MSG_ERR_LOGIN = "账号或密码错误！"
    MSG_ERR_DB_READ = "数据库读取失败: "
    MSG_ERR_FILE = "文件操作失败: "
    MSG_CONFIRM_DEL_BANK = "确定删除题库文件：{} 吗？"
    MSG_CONFIRM_DEL_QS = "确定要删除勾选的 {} 条题目吗？"
    MSG_TIP_SELECT = "请先勾选左侧方框"
    MSG_SUCCESS_IMPORT = "成功导入 {} 条题目"
    MSG_SUCCESS_EXPORT = "导出成功"

    HINT_SEARCH = "🔍 搜索内容..."
    HINT_STATUS = "💡 提示：点击文字选中 | 点击方框勾选 | Ctrl+滚轮缩放"


# *********************************************************************
# 类名：DatabaseManager
# 功能描述：数据库操作封装类 (DAO层)。
#           负责所有与 SQLite 交互的底层逻辑，确保连接的安全打开与关闭。
# *********************************************************************
class DatabaseManager:
    def __init__(self, db_path):
        # 初始化时传入数据库文件路径 (.db 文件)
        self.m_dbPath = db_path

    # *********************************************************************
    # 功能描述：通用执行 SQL 方法
    # 输入参数：query —— SQL 语句 (如 "SELECT * FROM table")
    #          params —— SQL 参数元组 (防止 SQL 注入)
    #          fetch —— 是否需要返回查询结果 (SELECT 语句传 True, INSERT/UPDATE 传 False)
    # *********************************************************************
    def execute(self, query, params=(), fetch=False):
        conn = None
        res = None
        try:
            # 1. 建立连接：如果文件不存在会自动创建
            conn = sqlite3.connect(self.m_dbPath)
            cursor = conn.cursor()
            # 2. 执行语句：使用参数化查询，安全且高效
            cursor.execute(query, params)
            # 3. 获取数据：如果 fetch 为 True，拉取所有结果
            if fetch:
                res = cursor.fetchall()
            # 4. 提交事务：写入操作(CUD)必须 commit 才能生效
            conn.commit()
        except sqlite3.Error as e:
            # 5. 错误捕获：防止程序崩溃，并弹窗提示用户
            print(f"Database Error: {e}")
            messagebox.showerror("数据库错误", f"{StringResources.MSG_ERR_DB_READ}{e}")
        finally:
            # 6. 资源释放：无论成功失败，必须关闭连接，防止文件被锁定
            if conn:
                conn.close()
        return res


# *********************************************************************
# 类名：QuestionBankSystem
# 功能描述：主应用程序类 (UI层 + 业务逻辑层)。
#           继承自 ctk.CTk (主窗口)，管理所有界面切换和业务流程。
# *********************************************************************
class QuestionBankSystem(ctk.CTk):
    def __init__(self):
        super().__init__()
        # 1. 基础窗口设置
        self.title(StringResources.TITLE_MAIN)
        self.geometry("1280x850")  # 默认宽 x 高

        # 2. 成员变量初始化 (状态管理)
        self.m_currentDbPath = None  # 当前选中的题库路径
        self.m_currentFontSize = 16  # 默认字号
        self.m_isSortAsc = False  # 排序状态 (False=降序/最新在前)
        self.m_checkedIds = set()  # 使用 set 存储被勾选的 ID，避免重复，O(1)查找
        self.m_dbManager = None  # 数据库管理实例占位

        # 3. UI 样式配置 (Treeview 表格样式)
        self.m_style = ttk.Style()
        self.m_style.theme_use("clam")  # 使用 clam 主题以便自定义颜色
        # 自定义选中行颜色：背景蓝，文字白
        self.m_style.map("Treeview", background=[('selected', '#0078D7')], foreground=[('selected', 'white')])
        self.update_font_style()  # 初始化字体

        # 4. 主容器：用于页面切换 (登录页 -> 列表页 -> 编辑页)
        self.m_container = ctk.CTkFrame(self)
        self.m_container.pack(fill="both", expand=True)

        # 5. 全局事件绑定：Ctrl + 滚轮实现字体缩放
        self.bind("<Control-MouseWheel>", self.change_font_size)

        # 6. 启动流程：显示登录界面
        self.show_login_frame()

    # *********************************************************************
    # 功能描述：UI 交互 - 动态调整字体大小 (Ctrl+滚轮)
    # *********************************************************************
    def change_font_size(self, event, direction=None):
        # 防止弹窗存在时，主界面也跟着缩放
        if self.focus_get() and isinstance(self.focus_get().winfo_toplevel(), ctk.CTkToplevel):
            return
        # 计算方向：向上滚(+1) 或 向下滚(-1)
        delta = direction if direction else (1 if event.delta > 0 else -1)
        new_size = self.m_currentFontSize + delta
        # 限制字号范围 (8-30)
        if 8 <= new_size <= 30:
            self.m_currentFontSize = new_size
            self.update_font_style()  # 更新样式配置
            # 如果表格存在，刷新表格行高
            if hasattr(self, 'm_tree') and self.m_tree.winfo_exists():
                self.refresh_table(keep_selection=True)

    # 功能描述：配置 Treeview 样式 (表头和内容)
    def update_font_style(self):
        # 行高随字号动态调整 (字号 * 2.5)
        row_height = int(self.m_currentFontSize * 2.5)
        # 设置表头字体 (加粗)
        self.m_style.configure("Treeview.Heading", font=("Microsoft YaHei UI", self.m_currentFontSize + 1, "bold"),
                               rowheight=row_height)
        # 设置正文字体
        self.m_style.configure("Treeview", font=("Microsoft YaHei UI", self.m_currentFontSize), rowheight=row_height)

    # ================= 1. 登录模块 =================
    def show_login_frame(self):
        self.clear_frame()  # 清空当前页面
        # 创建居中的登录卡片
        frame = ctk.CTkFrame(self.m_container)
        frame.place(relx=0.5, rely=0.5, anchor="center")

        ctk.CTkLabel(frame, text=StringResources.TITLE_LOGIN, font=("Microsoft YaHei UI", 24, "bold")).pack(pady=20,
                                                                                                            padx=60)

        self.m_entryUser = ctk.CTkEntry(frame, placeholder_text="账号", width=250)
        self.m_entryUser.pack(pady=10)
        # show="*" 用于隐藏密码输入
        self.m_entryPass = ctk.CTkEntry(frame, placeholder_text="密码", show="*", width=250)
        self.m_entryPass.pack(pady=10)

        ctk.CTkButton(frame, text=StringResources.BTN_LOGIN, command=self.verify_login, width=250).pack(pady=20)

    def verify_login(self):
        # 获取输入并去除首尾空格
        user = self.m_entryUser.get().strip()
        pwd = self.m_entryPass.get().strip()
        # 硬编码验证 (实际项目应查库)
        if user == "admin" and pwd == "123456":
            self.show_bank_selection_frame()  # 登录成功，跳转
        else:
            messagebox.showerror("错误", StringResources.MSG_ERR_LOGIN)

    # ================= 2. 题库选择模块 (文件管理) =================
    def show_bank_selection_frame(self):
        self.clear_frame()

        # --- 左侧边栏 ---
        sidebar = ctk.CTkFrame(self.m_container, width=220, corner_radius=0)
        sidebar.pack(side="left", fill="y")
        ctk.CTkLabel(sidebar, text="题库中心", font=("Microsoft YaHei UI", 22, "bold")).pack(pady=30)
        # 退出按钮：红色系
        ctk.CTkButton(sidebar, text=StringResources.BTN_EXIT, fg_color="#C62828", hover_color="#B71C1C",
                      text_color="white", command=self.show_login_frame).pack(side="bottom", pady=20)

        # --- 右侧主区域 ---
        main_area = ctk.CTkFrame(self.m_container)
        main_area.pack(side="right", fill="both", expand=True, padx=20, pady=20)
        ctk.CTkLabel(main_area, text="📂 我的题库列表", font=("Microsoft YaHei UI", 18)).pack(anchor="w", pady=10)

        # 工具栏
        tool_frame = ctk.CTkFrame(main_area, fg_color="transparent")
        tool_frame.pack(fill="x", pady=10)
        ctk.CTkButton(tool_frame, text=StringResources.BTN_ADD_BANK, command=self.create_new_bank).pack(side="left")
        ctk.CTkButton(tool_frame, text=StringResources.BTN_DEL_BANK, command=self.delete_bank, fg_color="#D32F2F",
                      hover_color="#B71C1C").pack(side="left", padx=10)

        # 题库文件列表 (Treeview)
        columns = ("name", "path")
        self.m_bankTree = ttk.Treeview(main_area, columns=columns, show="headings", height=15)
        self.m_bankTree.heading("name", text="题库名称")
        self.m_bankTree.column("name", width=200)
        self.m_bankTree.heading("path", text="文件路径")
        self.m_bankTree.column("path", width=500)
        self.m_bankTree.pack(fill="both", expand=True)
        # 绑定双击事件 -> 打开题库
        self.m_bankTree.bind("<Double-1>", self.on_bank_double_click)
        self.load_bank_list()  # 加载文件列表

    # 读取当前目录下的 .db 文件
    def load_bank_list(self):
        for item in self.m_bankTree.get_children():
            self.m_bankTree.delete(item)
        try:
            files = [f for f in os.listdir('.') if f.endswith('.db')]
            for f in files:
                # 显示文件名和绝对路径
                self.m_bankTree.insert("", "end", values=(f, os.path.abspath(f)))
        except Exception as e:
            messagebox.showerror("错误", f"{StringResources.MSG_ERR_FILE}{e}")

    # 创建新题库 (新建 SQLite 文件并建表)
    def create_new_bank(self):
        dialog = ctk.CTkInputDialog(text="请输入新题库名称:", title="新建")
        name = dialog.get_input()
        if name:
            if not name.endswith(".db"): name += ".db"
            # 初始化数据库连接并创建表结构
            temp_db = DatabaseManager(name)
            temp_db.execute('''CREATE TABLE IF NOT EXISTS questions (
                                id INTEGER PRIMARY KEY AUTOINCREMENT, 
                                type TEXT, 
                                content TEXT, 
                                options TEXT, 
                                answer TEXT)''')
            self.load_bank_list()  # 刷新列表

    # 删除题库文件
    def delete_bank(self):
        selected = self.m_bankTree.selection()
        if not selected: return
        filename = self.m_bankTree.item(selected[0], "values")[0]
        if messagebox.askyesno("确认", StringResources.MSG_CONFIRM_DEL_BANK.format(filename)):
            try:
                os.remove(filename)  # 物理删除文件
                self.load_bank_list()
            except Exception as e:
                messagebox.showerror("错误", f"{StringResources.MSG_ERR_FILE}{e}")

    # 双击进入题库编辑页面
    def on_bank_double_click(self, event):
        selected = self.m_bankTree.selection()
        if not selected: return
        self.m_currentDbPath = self.m_bankTree.item(selected[0], "values")[0]
        self.m_dbManager = DatabaseManager(self.m_currentDbPath)  # 初始化当前 DB 管理器
        self.m_checkedIds.clear()  # 清空上一轮的选中状态
        self.show_editor_frame()

    # ================= 3. 题目编辑模块 (核心 CRUD) =================
    def show_editor_frame(self):
        self.clear_frame()

        # --- 顶部导航栏 ---
        nav = ctk.CTkFrame(self.m_container, height=50)
        nav.pack(fill="x")
        # 返回按钮
        ctk.CTkButton(nav, text="< 返回列表", command=self.show_bank_selection_frame, width=100, fg_color="transparent",
                      border_width=1).pack(side="left", padx=10, pady=10)
        ctk.CTkLabel(nav, text=f"正在编辑: {self.m_currentDbPath}", font=("Microsoft YaHei UI", 16, "bold")).pack(
            side="left", padx=20)
        ctk.CTkLabel(nav, text=StringResources.HINT_STATUS, text_color="gray").pack(side="right", padx=20)

        # --- 操作工具栏 (功能按钮区) ---
        toolbar = ctk.CTkFrame(self.m_container)
        toolbar.pack(fill="x", padx=10, pady=5)

        # 左侧：增删、统计、排序
        ctk.CTkButton(toolbar, text=StringResources.BTN_ADD_Q, width=80, command=self.open_add_window).pack(side="left",
                                                                                                            padx=5)
        ctk.CTkButton(toolbar, text=StringResources.BTN_DEL_CHECKED, width=100, fg_color="#D32F2F",
                      hover_color="#B71C1C", command=self.delete_checked_questions).pack(side="left", padx=5)
        ctk.CTkButton(toolbar, text=StringResources.BTN_STATS, width=100, fg_color="#7B1FA2", hover_color="#4A148C",
                      command=self.show_statistics).pack(side="left", padx=5)
        # 排序按钮
        self.m_btnSort = ctk.CTkButton(toolbar, text="⇅ 排序: ID降序", width=120, command=self.toggle_sort,
                                       fg_color="#F57F17", hover_color="#F9A825")
        self.m_btnSort.pack(side="left", padx=5)

        # 右侧：导入导出
        ctk.CTkButton(toolbar, text=StringResources.BTN_IMPORT, width=80, command=self.import_data_step1).pack(
            side="right", padx=5)
        ctk.CTkButton(toolbar, text=StringResources.BTN_EXPORT_ALL, width=100, command=self.export_all_word).pack(
            side="right", padx=5)
        ctk.CTkButton(toolbar, text=StringResources.BTN_EXPORT_CHECKED, width=100, command=self.export_checked_word,
                      fg_color="#00796B", hover_color="#004D40").pack(side="right", padx=5)

        # --- 搜索栏 ---
        search_frame = ctk.CTkFrame(self.m_container, fg_color="transparent")
        search_frame.pack(fill="x", padx=10, pady=5)
        self.m_searchEntry = ctk.CTkEntry(search_frame, placeholder_text=StringResources.HINT_SEARCH, width=300)
        self.m_searchEntry.pack(side="left")
        ctk.CTkButton(search_frame, text=StringResources.BTN_SEARCH, width=60, command=self.refresh_table).pack(
            side="left", padx=5)
        ctk.CTkButton(search_frame, text=StringResources.BTN_RESET, width=60, fg_color="gray",
                      command=self.reset_search).pack(side="left", padx=5)

        # --- 数据展示表格 (Treeview) ---
        columns = ("Select", "ID", "题型", "题目内容", "选项", "答案")
        self.m_tree = ttk.Treeview(self.m_container, columns=columns, show="headings", selectmode="extended")

        # 定义表头
        self.m_tree.heading("Select", text="☑")  # 勾选列
        self.m_tree.column("Select", width=50, anchor="center")
        self.m_tree.heading("ID", text="ID")
        self.m_tree.column("ID", width=60, anchor="center")
        self.m_tree.heading("题型", text="题型")
        self.m_tree.column("题型", width=100, anchor="center")
        self.m_tree.heading("题目内容", text="题目内容")
        self.m_tree.column("题目内容", width=500)
        self.m_tree.heading("选项", text="选项")
        self.m_tree.column("选项", width=250)
        self.m_tree.heading("答案", text="答案")
        self.m_tree.column("答案", width=100)

        # 滚动条
        vsb = ttk.Scrollbar(self.m_container, orient="vertical", command=self.m_tree.yview)
        self.m_tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self.m_tree.pack(fill="both", expand=True, padx=10, pady=5)

        # 绑定事件：单击切换勾选，双击编辑
        self.m_tree.bind("<Button-1>", self.on_tree_click)
        self.m_tree.bind("<Double-1>", self.on_tree_double_click)

        self.refresh_table()  # 初始加载数据

    # 刷新表格数据 (包含 搜索 + 排序 + 勾选状态回填)
    def refresh_table(self, keep_selection=False):
        for item in self.m_tree.get_children():
            self.m_tree.delete(item)

        sql = "SELECT * FROM questions"
        params = ()
        keyword = self.m_searchEntry.get().strip()

        # 构建搜索条件
        if keyword:
            sql += " WHERE content LIKE ?"
            params = (f"%{keyword}%",)  # 模糊查询

        # 构建排序条件
        sql += f" ORDER BY id {'ASC' if self.m_isSortAsc else 'DESC'}"

        data = self.m_dbManager.execute(sql, params, fetch=True)
        if data:
            for row in data:
                q_id = row[0]
                # 根据 m_checkedIds 集合判断该行是否应显示勾选框
                check_mark = "☑" if q_id in self.m_checkedIds else "☐"
                content = row[2].replace('\n', ' ') if row[2] else ""  # 替换换行符，防止表格显示错乱
                self.m_tree.insert("", "end", values=(check_mark, row[0], row[1], content, row[3], row[4]))

    # 切换 ID 排序顺序
    def toggle_sort(self):
        self.m_isSortAsc = not self.m_isSortAsc
        self.m_btnSort.configure(text=f"⇅ 排序: ID{'升序' if self.m_isSortAsc else '降序'}")
        self.refresh_table(keep_selection=True)

    # 重置搜索
    def reset_search(self):
        self.m_searchEntry.delete(0, "end")
        self.refresh_table()

    # 表格单击事件处理 (实现自定义 Checkbox 逻辑)
    def on_tree_click(self, event):
        # 判断点击区域
        region = self.m_tree.identify("region", event.x, event.y)

        # 1. 点击表头 -> 全选/反选
        if region == "heading":
            col = self.m_tree.identify_column(event.x)
            if col == "#1": self.toggle_select_all()
            return None

        # 2. 点击单元格
        if region == "cell":
            col = self.m_tree.identify_column(event.x)
            row_id = self.m_tree.identify_row(event.y)
            # 如果点击的是第一列 (Select 列)
            if col == "#1":
                values = self.m_tree.item(row_id, "values")
                q_id = int(values[1])
                # 状态反转逻辑
                if q_id in self.m_checkedIds:
                    self.m_checkedIds.remove(q_id)
                    new_mark = "☐"
                else:
                    self.m_checkedIds.add(q_id)
                    new_mark = "☑"
                # 只更新显示，不需重新查库，提高性能
                new_vals = (new_mark,) + values[1:]
                self.m_tree.item(row_id, values=new_vals)
                return "break"  # 阻止事件冒泡，防止 Treeview 默认的行选中行为覆盖逻辑
            return None
        return None

    # 全选/取消全选逻辑
    def toggle_select_all(self):
        all_items = self.m_tree.get_children()
        if not all_items: return
        all_ids = [int(self.m_tree.item(item, "values")[1]) for item in all_items]

        # 如果当前页所有项都已被选，则执行“取消全选”
        if all(qid in self.m_checkedIds for qid in all_ids):
            for qid in all_ids: self.m_checkedIds.discard(qid)
        else:
            # 否则执行“全选”
            for qid in all_ids: self.m_checkedIds.add(qid)
        self.refresh_table(keep_selection=True)

    # 打开新增窗口
    def open_add_window(self):
        self.edit_window(is_add=True)

    # 双击编辑
    def on_tree_double_click(self, event):
        item = self.m_tree.selection()
        if not item:
            item = self.m_tree.identify_row(event.y)
        if not item: return
        item = item[0] if isinstance(item, tuple) else item
        vals = self.m_tree.item(item, "values")
        q_id = int(vals[1])
        # 编辑前重新查库，确保数据最新
        data = self.m_dbManager.execute("SELECT * FROM questions WHERE id=?", (q_id,), fetch=True)
        if data:
            self.edit_window(is_add=False, data=data[0])

    # *********************************************************************
    # 核心算法：自动题型推断
    # 逻辑：基于内容特征和答案格式进行启发式判断
    # *********************************************************************
    def auto_infer(self, content, options, answer):
        c = content.strip() if content else ""
        o = options.strip() if options else ""
        a = answer.strip().upper() if answer else ""

        # 1. 优先判断判断题 (特征：答案是 T/F/对/错)
        if a in ['T', 'F', 'TRUE', 'FALSE', 'YES', 'NO', '对', '错', '√', '×', '正确', '错误']:
            return "判断题"

        # 2. 推断选择题
        # 清洗答案，只保留字母 (如 "A, B" -> "AB")
        clean_ans = re.sub(r'[^A-Z]', '', a)
        # 检查题目中是否有 "(A)" 或 "A." 这种选项标记
        has_option_pattern = bool(re.search(r'(?:^|\s)[(（\[]?[A-G][)）\].、]', c))
        # 检查答案是否由纯字母组成 (防止填空题答案也是字母的情况)
        is_valid_choice_key = False
        if clean_ans:
            if re.fullmatch(r'[A-H]+', clean_ans) and len(clean_ans) < 10:
                is_valid_choice_key = True

        # 综合判断：有选项内容 OR 题目有选项格式 OR 答案像选项
        if (bool(o) or has_option_pattern or is_valid_choice_key) and is_valid_choice_key:
            if len(clean_ans) > 1:
                return "多选题"  # 答案由多个字母组成 -> 多选
            elif len(clean_ans) == 1:
                return "单选题"  # 答案只有一个字母 -> 单选

        # 3. 推断填空题 (特征：题目中有下划线、括号占位符)
        if re.search(r'_{2,}', c) or re.search(r'（\s*）', c) or re.search(r'\(\s*\)', c):
            return "填空题"

        return "主观题"  # 兜底逻辑

    # 题型名称标准化 (将 "Single Choice", "单选" 等统一为 "单选题")
    def normalize_type(self, raw_type):
        if not raw_type: return "未知题型"
        raw = raw_type.lower().strip()

        if "多选" in raw: return "多选题"
        if "单选" in raw: return "单选题"
        if "判断" in raw: return "判断题"
        if "填空" in raw: return "填空题"
        if "主观" in raw: return "主观题"

        if "multiple" in raw: return "多选题"
        if "single" in raw: return "单选题"
        if "true" in raw or "false" in raw or "judgment" in raw: return "判断题"
        if "fill" in raw or "completion" in raw: return "填空题"

        return "主观题"

    # 新增/编辑 通用弹窗
    def edit_window(self, is_add, data=None):
        win = ctk.CTkToplevel(self)
        win.title("题目编辑")
        win.geometry("600x600")
        win.grab_set()  # 模态窗口，禁止操作主界面

        # UI 布局：文本框、输入框、下拉框
        ctk.CTkLabel(win, text="内容:").pack(pady=5)
        txt = ctk.CTkTextbox(win, height=120)
        txt.pack(padx=20, fill="x")

        ctk.CTkLabel(win, text="选项:").pack(pady=5)
        ent_opt = ctk.CTkEntry(win)
        ent_opt.pack(padx=20, fill="x")

        ctk.CTkLabel(win, text="答案:").pack(pady=5)
        ent_ans = ctk.CTkEntry(win)
        ent_ans.pack(padx=20, fill="x")

        ctk.CTkLabel(win, text="题型:").pack(pady=5)
        var_type = ctk.StringVar(value="自动推断")
        # 支持手动选择或自动推断
        ctk.CTkComboBox(win, values=["自动推断", "单选题", "多选题", "判断题", "填空题", "主观题"],
                        variable=var_type).pack(pady=5)

        # 如果是编辑模式，回填数据
        if not is_add and data:
            txt.insert("1.0", data[2] if data[2] else "")
            ent_opt.insert(0, data[3] if data[3] else "")
            ent_ans.insert(0, data[4] if data[4] else "")
            var_type.set(data[1] if data[1] else "自动推断")

        # 保存回调
        def save():
            c = txt.get("1.0", "end").strip()
            o = ent_opt.get().strip()
            a = ent_ans.get().strip()
            t = var_type.get()
            if not c: return

            # 执行自动推断
            if t == "自动推断":
                t = self.auto_infer(c, o, a)
            t = self.normalize_type(t)

            if is_add:
                # 查重逻辑
                if self.m_dbManager.execute("SELECT id FROM questions WHERE content=?", (c,), fetch=True):
                    messagebox.showwarning("重复", "题目已存在")
                    return
                self.m_dbManager.execute("INSERT INTO questions (type,content,options,answer) VALUES (?,?,?,?)",
                                         (t, c, o, a))
            else:
                self.m_dbManager.execute("UPDATE questions SET type=?,content=?,options=?,answer=? WHERE id=?",
                                         (t, c, o, a, data[0]))
            self.refresh_table(keep_selection=True)
            win.destroy()

        ctk.CTkButton(win, text="保存", command=save).pack(pady=20)

    # 批量删除
    def delete_checked_questions(self):
        if not self.m_checkedIds:
            messagebox.showwarning("提示", StringResources.MSG_TIP_SELECT)
            return
        if messagebox.askyesno("确认", StringResources.MSG_CONFIRM_DEL_QS.format(len(self.m_checkedIds))):
            for qid in self.m_checkedIds:
                self.m_dbManager.execute("DELETE FROM questions WHERE id=?", (qid,))
            self.m_checkedIds.clear()  # 删除后清空选中集
            self.refresh_table()

    # --- 导入/导出模块 ---
    def import_data_step1(self):
        fp = filedialog.askopenfilename(filetypes=[("文本", "*.txt")])
        if not fp: return
        try:
            with open(fp, 'r', encoding='utf-8') as f:
                lines = [l.strip() for l in f.readlines() if l.strip()]
            # 第一步：解析TXT，第二步：弹窗让用户确认
            self.show_import_selector(self.parse_txt(lines))
        except Exception as e:
            messagebox.showerror("错误", f"{StringResources.MSG_ERR_FILE}{e}")

    # *********************************************************************
    # 核心算法：TXT 文件解析
    # 逻辑：状态机模式。遍历每一行，识别题目开始、选项、答案。
    # *********************************************************************
    def parse_txt(self, lines):
        qs = []
        # 预编译正则提高性能
        ptn_type = re.compile(r'^\[(.*?)]$')  # 匹配 [单选题]
        ptn_start = re.compile(r'^[(（]?\d+[)）.:、．]\s*(.*)')  # 匹配 1. 题目
        ptn_ans_line = re.compile(r'^(?:Answer|Ans|Key|Correct|答案|正确答案|参考答案)\s*[:：]?\s*(.*)',
                                  re.I)  # 匹配 答案: A
        ptn_opt = re.compile(r'^[(（\[]?([A-Ha-h])[)）\].:、．]\s*(.*)', re.I)  # 匹配 A. 选项
        ptn_eol_ans = re.compile(r'(\s+|[\t\u3000]+)([A-Ha-h]+)\s*$')  # 匹配题目行末尾的答案 (如 "题目内容 (A)")

        curr_type, i = "自动推断", 0

        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # 1. 检查是否是题型标记行
            if ptn_type.match(line):
                curr_type = ptn_type.match(line).group(1)
                i += 1
                continue

            # 2. 检查是否是题目开始 (数字开头)
            if ptn_start.match(line):
                content = ptn_start.match(line).group(1)
                opts, ans = [], ""
                i += 1

                # 内层循环：收集该题目的后续行 (选项、答案或多行题目)
                while i < len(lines):
                    sub = lines[i].strip()
                    if not sub:
                        i += 1
                        continue
                    # 如果遇到新题型标记或新题目开始，跳出内循环
                    if ptn_type.match(sub) or ptn_start.match(sub):
                        break

                    # 检查是否是显式的答案行
                    ans_match = ptn_ans_line.match(sub)
                    if ans_match:
                        ans = ans_match.group(1).strip()
                        i += 1
                        continue

                    # 检查是否是选项行
                    if ptn_opt.match(sub):
                        opts.append(sub)
                    else:
                        # 既不是答案也不是选项 -> 视为题目内容的换行续写
                        content += "\n" + sub
                    i += 1

                # 如果没有找到显式答案，尝试从题目末尾提取
                if not ans:
                    content = content.strip()
                    eol_match = ptn_eol_ans.search(content)
                    if eol_match:
                        potential_ans = eol_match.group(2)

                        if len(potential_ans) <= 5:
                            ans = potential_ans
                            content = content[:eol_match.start()].strip()  # 从内容中移除答案部分

                # 处理当前收集到的题目
                ft = curr_type
                if ft == "自动推断":
                    ft = self.auto_infer(content, " ".join(opts), ans)
                ft = self.normalize_type(ft)
                qs.append({"type": ft, "content": content, "options": " ".join(opts), "answer": ans})
            else:
                i += 1
        return qs

    # 显示导入预览窗口 (允许用户勾选需要导入的题目)
    def show_import_selector(self, qs):
        top = ctk.CTkToplevel(self)
        top.title(StringResources.TITLE_PREVIEW)
        top.geometry("900x600")
        top.grab_set()

        top_frame = ctk.CTkFrame(top, fg_color="transparent")
        top_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(top_frame, text=f"解析到 {len(qs)} 条题目，请勾选需要导入的项：", text_color="gray").pack(
            side="left")

        tree_frame = ctk.CTkFrame(top)
        tree_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # 预览表格
        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        columns = ("check", "type", "content", "answer")
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings", selectmode="none", yscrollcommand=vsb.set)

        vsb.config(command=tree.yview)
        vsb.pack(side="right", fill="y")
        tree.pack(side="left", fill="both", expand=True)

        tree.heading("check", text="☑")
        tree.column("check", width=40, anchor="center")
        tree.heading("type", text="题型")
        tree.column("type", width=100, anchor="center")
        tree.heading("content", text="内容")
        tree.column("content", width=550)
        tree.heading("answer", text="答案")
        tree.column("answer", width=100, anchor="center")

        # 默认全选
        selected_indices = set(range(len(qs)))

        for i, q in enumerate(qs):
            preview_content = q['content'].replace('\n', ' ')
            tree.insert("", "end", iid=str(i), values=("☑", q['type'], preview_content, q['answer']))

        # 预览窗口的点击逻辑 (类似主界面的勾选)
        def on_click(event):
            region = tree.identify("region", event.x, event.y)
            if region == "heading":
                col = tree.identify_column(event.x)
                if col == "#1":
                    if len(selected_indices) == len(qs):
                        selected_indices.clear()
                        new_mark = "☐"
                    else:
                        selected_indices.update(range(len(qs)))
                        new_mark = "☑"
                    tree.heading("check", text=new_mark)
                    for i in range(len(qs)):
                        vals = tree.item(str(i), "values")
                        tree.item(str(i), values=(new_mark,) + vals[1:])
            elif region == "cell":
                col = tree.identify_column(event.x)
                row_id = tree.identify_row(event.y)
                if col == "#1":
                    idx = int(row_id)
                    vals = tree.item(row_id, "values")
                    if idx in selected_indices:
                        selected_indices.remove(idx)
                        tree.item(row_id, values=("☐",) + vals[1:])
                    else:
                        selected_indices.add(idx)
                        tree.item(row_id, values=("☑",) + vals[1:])

        tree.bind("<Button-1>", on_click)

        # 确认导入逻辑
        def confirm():
            if not selected_indices:
                messagebox.showwarning("提示", "未勾选任何题目")
                return
            cnt = 0
            for idx in selected_indices:
                q = qs[idx]
                ft = self.normalize_type(q['type'])
                # 防止重复导入相同题目
                if not self.m_dbManager.execute("SELECT id FROM questions WHERE content=?", (q['content'],), True):
                    self.m_dbManager.execute(
                        "INSERT INTO questions (type,content,options,answer) VALUES (?,?,?,?)",
                        (ft, q['content'], q['options'], q['answer'])
                    )
                    cnt += 1
            self.refresh_table()
            top.destroy()
            messagebox.showinfo("完成", StringResources.MSG_SUCCESS_IMPORT.format(cnt))

        btn_frame = ctk.CTkFrame(top, fg_color="transparent")
        btn_frame.pack(pady=10)
        ctk.CTkButton(btn_frame, text="确认导入勾选题目", command=confirm, width=200).pack()

    # 导出全部为 Word
    def export_all_word(self):
        self._export_core(self.m_dbManager.execute("SELECT * FROM questions ORDER BY id ASC", fetch=True), "全部导出")

    # 导出选中为 Word
    def export_checked_word(self):
        if not self.m_checkedIds: return messagebox.showwarning("提示", StringResources.MSG_TIP_SELECT)
        ids = sorted(list(self.m_checkedIds))
        qs = []
        for i in ids: qs.append(self.m_dbManager.execute("SELECT * FROM questions WHERE id=?", (i,), fetch=True)[0])
        self._export_core(qs, "选中导出")
        return None

    # 导出核心逻辑 (生成 docx)
    def _export_core(self, qs, title):
        fp = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")])
        if not fp: return
        doc = Document()
        doc.add_heading(f'{title}: {os.path.basename(self.m_currentDbPath)}', 0)
        for i, q in enumerate(qs, 1):
            # 格式：1. [单选题] 题目内容 (加粗)
            doc.add_paragraph().add_run(f"{i}. [{q[1]}] {q[2]}").bold = True
            if q[3]: doc.add_paragraph(q[3])  # 选项
            doc.add_paragraph(f"答案: {q[4]}")
            doc.add_paragraph("-" * 20)  # 分隔线
        doc.save(fp)
        messagebox.showinfo("完成", StringResources.MSG_SUCCESS_EXPORT)

    # 统计图表展示 (Matplotlib)
    def show_statistics(self):
        win = ctk.CTkToplevel(self)
        win.title(StringResources.TITLE_STATS)
        win.geometry("800x600")

        # 获取数据
        raw = self.m_dbManager.execute("SELECT type FROM questions", fetch=True)
        if not raw: return

        # 计数
        counts_dict = defaultdict(int)
        for r in raw: counts_dict[self.normalize_type(r[0])] += 1

        # 配置字体支持中文
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['font.size'] = 12

        # 绘制饼图
        fig, ax = plt.subplots(figsize=(6, 4), dpi=100)
        w, t, a = ax.pie(list(counts_dict.values()), autopct='%1.1f%%', startangle=90, pctdistance=0.8,
                         textprops={'fontsize': 12, 'weight': 'bold'})
        for at in a: at.set_color('white')  # 百分比文字白色
        ax.legend(w, list(counts_dict.keys()), loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

        # 嵌入 Tkinter
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.get_tk_widget().pack(fill="both", expand=True, padx=20, pady=20)
        canvas.draw()

        # --- 以下为图表交互逻辑 (缩放和平移) ---
        def on_scroll(event):
            # 滚轮缩放
            scale = 1 / 1.1 if event.button == 'up' else 1.1
            if event.xdata is None: return
            xlim, ylim = ax.get_xlim(), ax.get_ylim()
            new_w, new_h = (xlim[1] - xlim[0]) * scale, (ylim[1] - ylim[0]) * scale
            relx = (xlim[1] - event.xdata) / (xlim[1] - xlim[0])
            rely = (ylim[1] - event.ydata) / (ylim[1] - ylim[0])
            ax.set_xlim([event.xdata - new_w * (1 - relx), event.xdata + new_w * relx])
            ax.set_ylim([event.ydata - new_h * (1 - rely), event.ydata + new_h * rely])
            canvas.draw()

        # 平移状态变量
        pan = {'p': False, 'x': 0, 'y': 0, 'xl': None, 'yl': None}

        def on_press(e):
            # 鼠标左键按下，记录起始位置
            if e.button == 1: pan.update({'p': True, 'x': e.x, 'y': e.y, 'xl': ax.get_xlim(), 'yl': ax.get_ylim()})

        def on_release(e):
            pan['p'] = False

        def on_motion(e):
            # 鼠标拖动，更新坐标轴范围
            if pan['p'] and e.inaxes:
                inv = ax.transData.inverted()
                d = inv.transform((e.x - pan['x'], e.y - pan['y'])) - inv.transform((0, 0))
                ax.set_xlim(pan['xl'] - d[0])
                ax.set_ylim(pan['yl'] - d[1])
                canvas.draw()

        # 绑定 Matplotlib 事件
        canvas.mpl_connect('scroll_event', on_scroll)
        canvas.mpl_connect('button_press_event', on_press)
        canvas.mpl_connect('button_release_event', on_release)
        canvas.mpl_connect('motion_notify_event', on_motion)

    # 清空主容器工具方法 (用于页面切换)
    def clear_frame(self):
        for w in self.m_container.winfo_children(): w.destroy()


if __name__ == "__main__":
    app = QuestionBankSystem()
    app.mainloop()